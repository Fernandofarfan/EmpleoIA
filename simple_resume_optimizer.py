import os
import logging
import re
from dotenv import load_dotenv
from MASTER_RESUME_PROMPT import get_formatted_prompt, USER_EXPERIENCE_CONTEXT

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class simpleResumeOptimizer:
    def __init__(self):
        self.api_key = os.getenv('GEMINI_API_KEY')
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        
        try:
            import google.generativeai as genai
            genai.configure(api_key=self.api_key)
            try:
                self.model = genai.GenerativeModel('gemini-1.5-pro')
                logger.info("Gemini Pro model loaded")
            except:
                self.model = genai.GenerativeModel('gemini-1.5-flash')
                logger.info("Gemini Flash model loaded")
        except ImportError:
            logger.error("google-generativeai package not installed")
            raise
    
    def optimize_resume(self, resume_text: str, job_description: str, job_title: str = "", company: str = "") -> str:
        try:
            prompt = get_formatted_prompt(
                job_description=job_description,
                base_resume_text=resume_text,
                user_context=USER_EXPERIENCE_CONTEXT
            )
            
            logger.info(f"Optimizing resume for {job_title} at {company}")
            
            response = self.model.generate_content(prompt)
            
            if response.text:
                optimized_text = response.text.strip()
                optimized_text = self._clean_and_format_resume(optimized_text)
                logger.info("Resume optimized successfully")
                return optimized_text
            else:
                logger.error("Empty response from Gemini")
                return resume_text
                
        except Exception as e:
            logger.error(f"Error optimizing resume: {e}")
            return resume_text
    
    def save_resume(self, resume_text: str, job_title: str, company: str, output_dir: str = "temp/resumes", format_type: str = "docx") -> str:
        os.makedirs(output_dir, exist_ok=True)
        
        safe_title = "".join(c for c in job_title if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
        safe_company = "".join(c for c in company if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
        
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format_type == "docx":
            filename = f"resume_{safe_title}_{safe_company}_{timestamp}.docx"
            filename = filename.replace(" ", "_")
            filepath = os.path.join(output_dir, filename)
            self._save_as_docx(resume_text, filepath)
        else:
            filename = f"resume_{safe_title}_{safe_company}_{timestamp}.txt"
            filename = filename.replace(" ", "_")
            filepath = os.path.join(output_dir, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(resume_text)
        
        logger.info(f"Resume saved to {filepath}")
        return filename
    
    def _save_as_docx(self, resume_text: str, filepath: str):
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            lines = resume_text.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                if any(keyword in line.upper() for keyword in ['PROFESSIONAL SUMMARY', 'TECHNICAL SKILLS', 'EXPERIENCE', 'PROJECTS', 'EDUCATION']):
                    if i > 0:
                        doc.add_paragraph()
                    p = doc.add_paragraph(line.upper())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.size = Pt(14)
                        run.font.bold = True
                    p.space_after = Pt(6)
                    i += 1
                    continue
                
                if '|' in line and len(line) < 200:
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                    p.space_after = Pt(3)
                    i += 1
                    continue
                
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    clean_line = line.lstrip('•-*').strip()
                    p = doc.add_paragraph(clean_line, style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(10)
                    p.space_after = Pt(2)
                    i += 1
                    continue
                
                if len(line) < 80 and not any(punct in line for punct in ['.', ',', ':', ';']) and i < 5:
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                    p.space_after = Pt(3)
                    i += 1
                    continue
                
                words = line.split()
                if len(words) > 0:
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(10)
                    p.space_after = Pt(2)
                
                i += 1
            
            doc.save(filepath)
            logger.info(f"DOCX resume saved successfully")
            
        except ImportError:
            logger.warning("python-docx not available, saving as text file")
            filepath_txt = filepath.replace('.docx', '.txt')
            with open(filepath_txt, 'w', encoding='utf-8') as f:
                f.write(resume_text)
        except Exception as e:
            logger.error(f"Error saving DOCX: {e}")
            filepath_txt = filepath.replace('.docx', '.txt')
            with open(filepath_txt, 'w', encoding='utf-8') as f:
                f.write(resume_text)
    
    def _parse_resume_text(self, text: str) -> dict:
        sections = {
            "header": [],
            "section_title": [],
            "content": [],
            "job_entry": []
        }
        
        lines = text.split('\n')
        current_section = "content"
        
        section_keywords = [
            "PROFESSIONAL SUMMARY", "SUMMARY",
            "TECHNICAL SKILLS", "SKILLS",
            "EXPERIENCE", "WORK EXPERIENCE",
            "PROJECTS", "PROJECT",
            "EDUCATION"
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            is_section_title = False
            for keyword in section_keywords:
                if keyword in line.upper() and len(line) < 50:
                    sections["section_title"].append(line)
                    current_section = "section_title"
                    is_section_title = True
                    break
            
            if not is_section_title:
                if '|' in line or '–' in line or (line.count('-') >= 2 and len(line) < 100):
                    sections["job_entry"].append(line)
                elif len(line) < 100 and not line.startswith('•') and not line.startswith('-'):
                    if current_section == "section_title":
                        sections["content"].append(line)
                    else:
                        sections["content"].append(line)
                else:
                    sections["content"].append(line)
        
        if not sections["header"]:
            first_lines = [l.strip() for l in lines[:3] if l.strip() and len(l.strip()) < 100]
            sections["header"] = first_lines[:2]
        
        return sections
    
    def _clean_and_format_resume(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        text = re.sub(r'([a-z])(\d)', r'\1 \2', text)
        text = re.sub(r'(\d)([A-Z])', r'\1 \2', text)
        text = re.sub(r'([a-z])([A-Z][a-z])', r'\1 \2', text)
        text = re.sub(r'for([a-z])', r'for \1', text)
        text = re.sub(r'in([A-Z])', r'in \1', text)
        text = re.sub(r'with([A-Z])', r'with \1', text)
        text = re.sub(r'on([A-Z])', r'on \1', text)
        text = re.sub(r'and([a-z])', r'and \1', text)
        text = re.sub(r'([a-z])([A-Z]{2,})', r'\1 \2', text)
        
        lines = text.split('\n')
        cleaned_lines = []
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_paragraph:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                if cleaned_lines and cleaned_lines[-1] != "":
                    cleaned_lines.append("")
                continue
            
            line = re.sub(r'\s+', ' ', line)
            
            section_keywords = ['PROFESSIONAL SUMMARY', 'TECHNICAL SKILLS', 'EXPERIENCE', 'PROJECTS', 'EDUCATION']
            is_section = any(keyword in line.upper() for keyword in section_keywords) and len(line) < 50
            
            if is_section:
                if current_paragraph:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                if cleaned_lines and cleaned_lines[-1] != "":
                    cleaned_lines.append("")
                cleaned_lines.append(line.upper())
                cleaned_lines.append("")
                continue
            
            if '|' in line and len(line) < 200:
                if current_paragraph:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                if cleaned_lines and cleaned_lines[-1] != "":
                    cleaned_lines.append("")
                cleaned_lines.append(line)
                continue
            
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                if current_paragraph:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                clean_line = line.lstrip('•-*').strip()
                cleaned_lines.append('• ' + clean_line)
                continue
            
            if len(line) < 80 and not any(punct in line for punct in ['.', ',', ':', ';']) and len(cleaned_lines) < 10:
                if current_paragraph:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                if cleaned_lines and cleaned_lines[-1] != "":
                    cleaned_lines.append("")
                cleaned_lines.append(line)
                continue
            
            if line.endswith('.') or line.endswith('.'):
                current_paragraph.append(line)
                if len(current_paragraph) > 0:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
            else:
                current_paragraph.append(line)
        
        if current_paragraph:
            cleaned_lines.append(' '.join(current_paragraph))
        
        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        result = re.sub(r' +', ' ', result)
        result = re.sub(r'([a-z])([A-Z][a-z])', r'\1 \2', result)
        
        return result.strip()

